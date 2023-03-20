
from docx import Document


import lib



doc = Document()


lib.add_heading('sdfsdf', doc)
lib.add_paragraph('sd333333333', doc)
lib.add_formula('Q_(от)=q_(0)*a*K*V_(a)*(t_(B)-t_(H))', doc)
doc.save('example.docx')

