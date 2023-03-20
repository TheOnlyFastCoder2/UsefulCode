from docx.shared import Pt, RGBColor
from asciitomathml.asciitomathml import AsciiMathML

import mathml2omml
from docx.oxml import parse_xml

def add_heading(text, doc):
    heading = doc.add_heading().add_run(text)
    heading.font.color.rgb = RGBColor(0, 0, 0)
    heading.font.name = 'Times New Roman'
    heading.font.size = Pt(14)
    return heading


def add_paragraph(text, doc):
    paragraph = doc.add_paragraph().add_run(text)
    paragraph.font.color.rgb = RGBColor(0, 0, 0)
    paragraph.font.name = 'Times New Roman'
    paragraph.font.size = Pt(12)
    return paragraph


def add_formula(str, doc):
    math_obj = AsciiMathML()
    math_obj.parse_string(str)
    mathml = math_obj.to_xml_string(encoding='unicode')

    omml = mathml2omml.convert(mathml)
    omml_xml = '<p xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">{}</p>'.format(omml)
    omml_el = parse_xml(omml_xml)[0]

    p = doc.add_paragraph()
    p._p.append(omml_el)

